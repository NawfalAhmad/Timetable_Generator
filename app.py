from flask import Flask, render_template, request, send_file
import random
from docx import Document
import os

app = Flask(__name__)

time_slots = [
    "8.00 A.M - 8.50 A.M",
    "8.50 A.M - 9.40 A.M",
    "BREAK (9.40 A.M - 10.10 A.M)",
    "10.10 A.M - 11.00 A.M",
    "11.00 A.M - 11.50 A.M",
    "11.50 A.M - 12.40 P.M",
    "LUNCH (12.40 P.M - 1.30 P.M)",
    "1.30 P.M - 2.15 P.M",
    "2.15 P.M - 3.00 P.M"
]

days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]

def create_docx(institution_info, timetable, file_path="timetable.docx"):
    doc = Document()
    # Header / Institution details
    doc.add_heading(institution_info.get('institution_name', ''), level=1)
    details = f"Department: {institution_info.get('department','')}\n"
    details += f"Calendar Year: {institution_info.get('calendar_year','')}\n"
    details += f"SEMESTER: {institution_info.get('semester','')}\n"
    details += f"Class: {institution_info.get('class_name','')}\n"
    details += f"Semester Number: {institution_info.get('semester_number','')}\n"
    details += f"Venue: {institution_info.get('venue','')}\n"
    doc.add_paragraph(details)

    for day in days:
        doc.add_heading(day, level=2)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Time'
        hdr_cells[1].text = 'Subject'
        for i in range(len(time_slots)):
            row_cells = table.add_row().cells
            row_cells[0].text = time_slots[i]
            row_cells[1].text = timetable[day][i]
        doc.add_paragraph('')

    doc.save(file_path)
    return file_path

@app.route('/')
def index():
    # default values to show in form
    defaults = {
        "institution_name": "RAJALAKSHMI INSTITUTE OF TECHNOLOGY, An Autonomous Institution, Affiliated to Anna University, Chennai, Kuthambakkam Post, Chennai - 600 124.",
        "department": "",
        "calendar_year": "",
        "semester": "ODD",
        "class_name": "",
        "semester_number": "",
        "venue": ""
    }
    return render_template('index.html', defaults=defaults)

@app.route('/generate', methods=['POST'])
def generate():
    try:
        num_subjects = int(request.form.get('num_subjects', 1))
    except:
        num_subjects = 1

    subjects = []
    for i in range(num_subjects):
        s = request.form.get(f'subject{i+1}', '').strip()
        if s == "":
            s = f"Subject-{i+1}"
        subjects.append(s)

    # Collect institution details
    institution_info = {
        'institution_name': request.form.get('institution_name', '').strip(),
        'department': request.form.get('department', '').strip(),
        'calendar_year': request.form.get('calendar_year', '').strip(),
        'semester': request.form.get('semester', '').strip(),
        'class_name': request.form.get('class_name', '').strip(),
        'semester_number': request.form.get('semester_number', '').strip(),
        'venue': request.form.get('venue', '').strip()
    }

    # Build randomized timetable
    timetable = {}
    for day in days:
        slots = []
        for slot in time_slots:
            if "BREAK" in slot or "LUNCH" in slot:
                slots.append(slot)
            else:
                slots.append(random.choice(subjects))
        timetable[day] = slots

    # Create .docx file (overwrite each time)
    file_path = "timetable.docx"
    create_docx(institution_info, timetable, file_path=file_path)

    return render_template('timetable.html', timetable=timetable, days=days, time_slots=time_slots, institution=institution_info)

@app.route('/download')
def download():
    file_path = "timetable.docx"
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    else:
        return "File not found", 404

if __name__ == "__main__":
    app.run(debug=True)
